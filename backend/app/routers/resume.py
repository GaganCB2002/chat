import json
import os
import re
import tempfile
import uuid
from datetime import datetime
from io import BytesIO
from typing import Optional

import httpx
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import Response
from fpdf import FPDF

router = APIRouter(prefix="/api/resume", tags=["resume"])

OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://127.0.0.1:11434")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY") or os.getenv("VITE_GEMINI_API_KEY", "")

MAX_FILE_SIZE = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx"}

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if doc.needs_pass:
            raise HTTPException(status_code=400, detail="Password-protected PDF. Please provide an unlocked file.")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from PDF. The file may be scanned or image-based. Please provide a text-based PDF.")
        return text.strip()
    except ImportError:
        raise HTTPException(status_code=500, detail="PDF parsing library not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)
        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from DOCX. The file may be empty or corrupted.")
        return text.strip()
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX parsing library not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {str(e)}")

def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}. Please upload PDF or DOCX.")

async def validate_file(file: UploadFile) -> bytes:
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported file type '{ext}'. Supported: PDF, DOCX")
    contents = await file.read()
    if not contents or len(contents) == 0:
        raise HTTPException(status_code=400, detail="Empty file")
    if len(contents) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large ({len(contents) / 1024 / 1024:.1f} MB). Max: 10 MB")
    return contents

OLLAMA_MODEL = os.getenv("OLLAMA_RESUME_MODEL", "qwen3.5:2b")

async def call_ollama(prompt: str, system: str = None, model: str = None) -> str:
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})
    use_model = model or OLLAMA_MODEL
    try:
        async with httpx.AsyncClient(timeout=180) as client:
            res = await client.post(f"{OLLAMA_HOST}/api/chat", json={
                "model": use_model,
                "messages": messages,
                "stream": False
            })
            if res.is_success:
                data = res.json()
                content = data.get("message", {}).get("content", "")
                if content:
                    return content
                raise HTTPException(status_code=502, detail="Ollama returned empty response")
            detail = res.text[:300]
            try:
                err_detail = res.json().get("error", detail)
            except Exception:
                err_detail = detail
            raise HTTPException(status_code=502, detail=f"Ollama error: {err_detail}")
    except httpx.ConnectError:
        raise HTTPException(status_code=503, detail="Ollama is not running. Please start Ollama and try again.")
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Ollama request timed out. The model may still be loading. Please try again.")

async def call_gemini(prompt: str, system: str = None) -> str:
    if not GEMINI_API_KEY:
        raise HTTPException(status_code=400, detail="Gemini API key not configured. Cannot use Gemini model.")
    full_prompt = f"{system}\n\n{prompt}" if system else prompt
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
    try:
        async with httpx.AsyncClient(timeout=120) as client:
            res = await client.post(
                url,
                json={"contents": [{"parts": [{"text": full_prompt}]}]},
                headers={
                    "Content-Type": "application/json",
                    "X-Goog-Api-Key": GEMINI_API_KEY
                }
            )
            if res.is_success:
                data = res.json()
                candidates = data.get("candidates", [])
                if candidates:
                    text = candidates[0].get("content", {}).get("parts", [{}])[0].get("text", "")
                    if text:
                        return text
                raise HTTPException(status_code=502, detail="Gemini returned empty response")
            detail = res.text[:500]
            try:
                err_data = res.json().get("error", {})
                err_detail = err_data.get("message", detail)
                err_code = err_data.get("code", res.status_code)
            except Exception:
                err_detail = detail
                err_code = res.status_code
            if res.status_code == 429 or err_code == 429:
                raise HTTPException(status_code=429, detail=f"Gemini quota exceeded: {err_detail}")
            raise HTTPException(status_code=502, detail=f"Gemini error ({res.status_code}): {err_detail}")
    except httpx.ConnectError:
        raise HTTPException(status_code=503, detail="Gemini API not reachable. Check your internet connection.")
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Gemini request timed out.")

async def ai_analyze(prompt: str, system: str = None, model: str = None) -> str:
    use_model = (model or "").lower()

    if use_model == "gemini" or "gemini" in use_model:
        if not GEMINI_API_KEY:
            raise HTTPException(status_code=400, detail="Gemini API key not configured. Cannot use Gemini model.")
        return await call_gemini(prompt, system)

    return await call_ollama(prompt, system, use_model or None)

@router.post("/parse-resume")
async def parse_resume(file: Optional[UploadFile] = File(None), text: Optional[str] = Form(None), model: Optional[str] = Form(None)):
    resume_text = ""
    if file:
        contents = await validate_file(file)
        resume_text = extract_text(contents, file.filename)
    elif text:
        resume_text = text.strip()
    else:
        raise HTTPException(status_code=400, detail="Provide a file or paste the resume text")
    
    if not resume_text:
        raise HTTPException(status_code=400, detail="Empty resume text")
        
    system = "You are a resume parsing expert. Extract structured information from resumes."
    prompt = f"""Parse the following resume and return a JSON object with these fields (use null for missing):
- name: Full name
- email: Email address
- phone: Phone number
- location: City/State
- linkedin: LinkedIn URL
- github: GitHub URL
- portfolio: Portfolio URL
- summary: Professional summary
- skills: Array of skill strings
- technical_skills: Object with categories (programming_languages, frameworks, databases, cloud_platforms, devops, ai_ml, tools) each as array of strings
- experience: Array of objects {{title, company, location, start_date, end_date, description, bullet_points: string[]}}
- projects: Array of objects {{name, description, technologies: string[], link}}
- education: Array of objects {{degree, institution, location, graduation_date, gpa}}
- certifications: Array of objects {{name, issuer, date}}
- achievements: Array of strings
- publications: Array of objects {{title, journal, date, link}}

Resume text:
{resume_text}

Return ONLY valid JSON, no markdown, no code blocks."""
    result = await ai_analyze(prompt, system, model)
    result = result.strip()
    if result.startswith("```"):
        result = re.sub(r'^```(?:json)?\s*', '', result)
        result = re.sub(r'\s*```$', '', result)
    try:
        parsed = json.loads(result)
    except json.JSONDecodeError:
        return {"raw_text": resume_text, "parsed": None, "error": "AI parsing failed"}
    return {"raw_text": resume_text, "parsed": parsed}

@router.post("/parse-jd")
async def parse_jd(file: Optional[UploadFile] = File(None), text: Optional[str] = Form(None), model: Optional[str] = Form(None)):
    jd_text = ""
    if file:
        contents = await validate_file(file)
        jd_text = extract_text(contents, file.filename)
    elif text:
        jd_text = text.strip()
    else:
        raise HTTPException(status_code=400, detail="Provide a file or paste the job description")
    if not jd_text:
        raise HTTPException(status_code=400, detail="Empty job description")
    system = "You are a job description parsing expert."
    prompt = f"""Parse the following job description and return a JSON object with these fields:
- job_title: string
- required_skills: array of strings
- preferred_skills: array of strings
- responsibilities: array of strings
- experience_requirements: string
- education: string
- certifications: array of strings
- keywords: array of strings
- technologies: array of strings
- soft_skills: array of strings
- summary: brief 2-3 sentence summary of the role

JD text:
{jd_text}

Return ONLY valid JSON, no markdown, no code blocks."""
    result = await ai_analyze(prompt, system, model)
    result = result.strip()
    if result.startswith("```"):
        result = re.sub(r'^```(?:json)?\s*', '', result)
        result = re.sub(r'\s*```$', '', result)
    try:
        parsed = json.loads(result)
    except json.JSONDecodeError:
        return {"raw_text": jd_text, "parsed": None, "error": "AI parsing failed"}
    return {"raw_text": jd_text, "parsed": parsed}

@router.post("/analyze")
async def analyze_resume(resume_text: str = Form(...), jd_text: str = Form(...), resume_json: str = Form("{}"), jd_json: str = Form("{}"), model: Optional[str] = Form(None)):
    system = "You are an ATS (Applicant Tracking System) expert. Analyze resumes against job descriptions. Be thorough and precise."
    prompt = f"""Compare this resume against the job description and return a JSON object with EXACTLY these fields:
    IMPORTANT: Be extremely thorough. Identify EVERY skill, technology, and tool mentioned in the JD. Compare each one against the resume.

1. ats_score: overall ATS compatibility score (0-100)
2. keyword_match: keyword match percentage (0-100)
3. skills_match: skills match percentage (0-100)
4. experience_match: experience match percentage (0-100)
5. education_match: education match percentage (0-100)
6. formatting_score: formatting score (0-100)

7. missing_skills: array of objects {{skill, category, relevance}} 
   - category is one of: programming_languages, frameworks, databases, cloud_platforms, devops, ai_ml, soft_skills, certifications, tools, other
   - relevance: "high", "medium", or "low" based on how important the skill is in the JD
   - INCLUDE EVERY skill from the JD that is not explicitly listed in the resume

8. suggestions: array of objects {{section, recommendation, reason}} 
   - section is one of: professional_summary, project_descriptions, experience_bullets, skills_section, achievements, formatting, ats_keywords
   - For skills_section: recommend adding specific missing skills

9. keyword_coverage: object with {{matched_keywords: string[], missing_keywords: string[], coverage_percentage: number}}
   - missing_keywords: include every JD keyword/technology/skill not found in resume

10. strengths: array of strings (what the resume does well)
11. weaknesses: array of strings (areas to improve, especially missing skills)

Resume JSON: {resume_json}
Resume Text: {resume_text[:4000]}

Job Description JSON: {jd_json}
Job Description Text: {jd_text[:4000]}

Return ONLY valid JSON, no markdown, no code blocks."""
    result = await ai_analyze(prompt, system, model)
    result = result.strip()
    if result.startswith("```"):
        result = re.sub(r'^```(?:json)?\s*', '', result)
        result = re.sub(r'\s*```$', '', result)
    try:
        return json.loads(result)
    except json.JSONDecodeError:
        return {
            "ats_score": 50, "keyword_match": 50, "skills_match": 50,
            "experience_match": 50, "education_match": 50, "formatting_score": 50,
            "missing_skills": [], "suggestions": [],
            "keyword_coverage": {"matched_keywords": [], "missing_keywords": [], "coverage_percentage": 50},
            "strengths": ["Resume submitted for analysis"], "weaknesses": ["AI analysis returned unexpected format"],
            "error": "Analysis parsing failed - AI returned invalid JSON"
        }

@router.post("/generate-optimized")
async def generate_optimized(
    resume_text: str = Form(...),
    jd_text: str = Form(...),
    resume_json: str = Form("{}"),
    jd_json: str = Form("{}"),
    analysis_json: str = Form("{}"),
    edits_json: str = Form("{}"),
    model: Optional[str] = Form(None)
):
    system = """You are a professional resume writer. Generate an optimized resume.

CRITICAL RULES:
1. Use ONLY factual information from the provided resume for experience, education, projects, certifications, and publications.
2. NEVER invent companies, degrees, certifications, projects, or years of experience.
3. HOWEVER, you MUST add all missing technical skills, tools, technologies, programming languages, frameworks, databases, and keywords from the Job Description into the resume's skills section.
4. Compare the JD's required/preferred skills, technologies, and keywords against the resume's current skills.
5. Add any skill from the JD that is not already in the resume to both the 'skills' array AND the appropriate category in 'technical_skills'.
6. Enhance existing bullet points with JD-relevant keywords while keeping them factual.
7. Use strong action verbs and quantify achievements where possible.
8. Keep descriptions concise and impactful."""

    prompt = f"""Generate an optimized resume as a JSON object with these sections:
- name: string
- email: string
- phone: string
- location: string
- linkedin: string
- github: string
- portfolio: string
- summary: string (3-4 lines, keyword-optimized, must reflect JD requirements)
- skills: array of strings (MUST include ALL skills from JD that are missing in resume)
- technical_skills: object with categories as arrays (MUST add missing JD skills to appropriate categories)
- experience: array of objects {{title, company, location, start_date, end_date, bullet_points: string[]}}
- projects: array of objects {{name, description, technologies, bullet_points: string[]}}
- education: array of objects {{degree, institution, location, graduation_date, gpa}}
- certifications: array of objects {{name, issuer, date}}
- achievements: array of strings
- publications: array of objects {{title, journal, date, link}}

Resume: {resume_text[:5000]}
Resume JSON: {resume_json}
Job Description: {jd_text[:5000]}
JD JSON: {jd_json}
ATS Analysis (shows missing skills): {analysis_json[:3000]}
User edits (apply these): {edits_json}

IMPORTANT: Identify all skills, technologies, and tools mentioned in the Job Description that are NOT in the Resume. Add ALL of them to the 'skills' array and 'technical_skills' object. This is critical for ATS optimization.

Return ONLY valid JSON, no markdown, no code blocks."""
    result = await ai_analyze(prompt, system, model)
    result = result.strip()
    if result.startswith("```"):
        result = re.sub(r'^```(?:json)?\s*', '', result)
        result = re.sub(r'\s*```$', '', result)
    try:
        data = json.loads(result)
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="AI generation returned invalid JSON")
    return data

def build_latex_classic(data: dict) -> str:
    name = escape_latex(data.get("name", "Your Name"))
    email = escape_latex(data.get("email", ""))
    phone = escape_latex(data.get("phone", ""))
    location = escape_latex(data.get("location", ""))
    linkedin = escape_latex(data.get("linkedin", ""))
    github = escape_latex(data.get("github", ""))
    portfolio = escape_latex(data.get("portfolio", ""))
    summary = escape_latex(data.get("summary", ""))

    contact_parts = []
    for c in [email, phone, location]:
        if c:
            contact_parts.append(c)
    contact_line = " $|$ ".join(contact_parts)
    links = " $|$ ".join(filter(None, [linkedin, github, portfolio]))

    tex = "\\documentclass[10pt,a4paper]{article}\n"
    tex += "\\usepackage[utf8]{inputenc}\n"
    tex += "\\usepackage[T1]{fontenc}\n"
    tex += "\\usepackage{geometry}\n"
    tex += "\\geometry{margin=0.7in}\n"
    tex += "\\usepackage{parskip}\n"
    tex += "\\usepackage{hyperref}\n"
    tex += "\\hypersetup{colorlinks=true,linkcolor=black,urlcolor=blue,pdfborder={0 0 0}}\n"
    tex += "\\usepackage{enumitem}\n"
    tex += "\\setlist{nosep,left=0pt,label=\\textbullet,itemsep=1pt}\n"
    tex += "\\renewcommand{\\familydefault}{\\sfdefault}\n"
    tex += "\\pagestyle{empty}\n"
    tex += "\\begin{document}\n"
    tex += "\\begin{center}\n"
    tex += "{\\Huge\\bfseries " + name + "}\\\\[6pt]\n"
    tex += contact_line + "\\\\\n"
    if links:
        tex += "\\small{" + links + "}\\\\\n"
    tex += "\\end{center}\n"
    tex += "\\vspace{-2pt}\n"
    tex += "\\rule{\\textwidth}{0.5pt}\n"
    tex += "\\vspace{4pt}\n"

    if summary:
        tex += section_block("Professional Summary", summary)

    skills = data.get("skills", [])
    if skills:
        tex += section_block("Skills", ", ".join(skills))

    tech_skills = data.get("technical_skills", {})
    if tech_skills:
        lines = []
        for cat, items in tech_skills.items():
            if items:
                label = cat.replace("_", " ").title()
                lines.append(f"\\textbf{{{label}}}: {', '.join(items)}")
        if lines:
            tex += "\\section*{Technical Skills}\n"
            tex += "\\vspace{-2pt}\\rule{\\textwidth}{0.5pt}\\vspace{4pt}\n"
            for l in lines:
                tex += l + "\\\\\n"
            tex += "\\vspace{4pt}\n"

    experience = data.get("experience", [])
    if experience:
        tex += "\\section*{Experience}\n"
        tex += "\\vspace{-2pt}\\rule{\\textwidth}{0.5pt}\\vspace{4pt}\n"
        for exp in experience:
            title = escape_latex(exp.get("title", ""))
            company = escape_latex(exp.get("company", ""))
            loc = escape_latex(exp.get("location", ""))
            sd = exp.get("start_date", "")
            ed = exp.get("end_date", "")
            tex += f"\\textbf{{{title}}} at \\textbf{{{company}}} \\hfill {sd} -- {ed}\\\\[2pt]\n"
            if loc:
                tex += f"\\textit{{{loc}}}\\\\[2pt]\n"
            bullets = exp.get("bullet_points", []) or [exp.get("description", "")]
            tex += "\\begin{itemize}[left=0pt]\n"
            for b in bullets:
                tex += "\\item " + escape_latex(b) + "\n"
            tex += "\\end{itemize}\n"
            tex += "\\vspace{2pt}\n"

    projects = data.get("projects", [])
    if projects:
        tex += "\\section*{Projects}\n"
        tex += "\\vspace{-2pt}\\rule{\\textwidth}{0.5pt}\\vspace{4pt}\n"
        for proj in projects:
            pname = escape_latex(proj.get("name", ""))
            techs = proj.get("technologies", [])
            tex += f"\\textbf{{{pname}}}"
            if techs:
                tex += f" \\hfill {{{', '.join(techs)}}}\\\\[2pt]\n"
            else:
                tex += "\\\\[2pt]\n"
            desc = escape_latex(proj.get("description", ""))
            bullets = proj.get("bullet_points", [])
            items = bullets if bullets else ([desc] if desc else [])
            if items:
                tex += "\\begin{itemize}[left=0pt]\n"
                for b in items:
                    tex += "\\item " + escape_latex(b) + "\n"
                tex += "\\end{itemize}\n"
            tex += "\\vspace{2pt}\n"

    education = data.get("education", [])
    if education:
        tex += "\\section*{Education}\n"
        tex += "\\vspace{-2pt}\\rule{\\textwidth}{0.5pt}\\vspace{4pt}\n"
        for edu in education:
            deg = escape_latex(edu.get("degree", ""))
            inst = escape_latex(edu.get("institution", ""))
            gd = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")
            tex += f"\\textbf{{{deg}}} -- {inst}"
            if gd:
                tex += f" \\hfill {gd}"
            tex += "\\\\[2pt]\n"
            if gpa:
                tex += f"GPA: {gpa}\\\\[2pt]\n"
            tex += "\\vspace{2pt}\n"

    certs = data.get("certifications", [])
    if certs:
        tex += "\\section*{Certifications}\n"
        tex += "\\vspace{-2pt}\\rule{\\textwidth}{0.5pt}\\vspace{4pt}\n"
        tex += "\\begin{itemize}[left=0pt]\n"
        for c in certs:
            cname = escape_latex(c.get("name", ""))
            issuer = escape_latex(c.get("issuer", ""))
            cd = c.get("date", "")
            parts2 = filter(None, [cname, issuer, cd])
            tex += "\\item " + " -- ".join(parts2) + "\n"
        tex += "\\end{itemize}\n"

    achievements = data.get("achievements", [])
    if achievements:
        tex += section_block("Achievements", [escape_latex(a) for a in achievements])

    publications = data.get("publications", [])
    if publications:
        tex += "\\section*{Publications}\n"
        tex += "\\vspace{-2pt}\\rule{\\textwidth}{0.5pt}\\vspace{4pt}\n"
        for pub in publications:
            ptitle = escape_latex(pub.get("title", ""))
            pjournal = escape_latex(pub.get("journal", ""))
            pdate = pub.get("date", "")
            plink = escape_latex(pub.get("link", ""))
            tex += "\\textbf{" + ptitle + "}"
            if pjournal:
                tex += " \\emph{" + pjournal + "}"
            if pdate:
                tex += " (" + pdate + ")"
            tex += "\\\\\n"
            if plink:
                tex += "\\url{" + plink + "}\\\\[4pt]\n"
            else:
                tex += "\\vspace{4pt}\n"

    tex += "\\end{document}\n"
    return tex


def build_latex_modern(data: dict) -> str:
    name = escape_latex(data.get("name", "Your Name"))
    email = escape_latex(data.get("email", ""))
    phone = escape_latex(data.get("phone", ""))
    location = escape_latex(data.get("location", ""))
    linkedin = escape_latex(data.get("linkedin", ""))
    github = escape_latex(data.get("github", ""))
    portfolio = escape_latex(data.get("portfolio", ""))
    summary = escape_latex(data.get("summary", ""))

    contact_parts = []
    for c in [email, phone, location]:
        if c:
            contact_parts.append(c)
    contact_line = " \\quad ".join(contact_parts)
    links = " \\quad ".join(filter(None, [linkedin, github, portfolio]))

    tex = "\\documentclass[10pt,a4paper]{article}\n"
    tex += "\\usepackage[utf8]{inputenc}\n"
    tex += "\\usepackage[T1]{fontenc}\n"
    tex += "\\usepackage{geometry}\n"
    tex += "\\geometry{margin=0.6in}\n"
    tex += "\\usepackage{parskip}\n"
    tex += "\\usepackage{hyperref}\n"
    tex += "\\hypersetup{colorlinks=true,linkcolor=teal,urlcolor=teal}\n"
    tex += "\\usepackage{enumitem}\n"
    tex += "\\setlist{nosep,left=0pt,label=\\textbullet,itemsep=0.5pt}\n"
    tex += "\\usepackage{xcolor}\n"
    tex += "\\definecolor{accent}{HTML}{0E7490}\n"
    tex += "\\renewcommand{\\familydefault}{\\sfdefault}\n"
    tex += "\\pagestyle{empty}\n"
    tex += "\\begin{document}\n"
    tex += "\\begin{center}\n"
    tex += "{\\huge\\bfseries\\color{accent} " + name + "}\\\\[4pt]\n"
    tex += "\\small{" + contact_line + "}\\\\\n"
    if links:
        tex += "\\footnotesize{" + links + "}\\\\\n"
    tex += "\\end{center}\n"
    tex += "\\vspace{2pt}\n"
    tex += "{\\color{accent}\\rule{\\textwidth}{1.5pt}}\n"
    tex += "\\vspace{6pt}\n"

    if summary:
        tex += "\\noindent{\\bfseries\\large\\color{accent}Professional Summary}\\\\[4pt]\n"
        tex += summary + "\\n\n\\vspace{6pt}\n"

    skills = data.get("skills", [])
    if skills:
        tex += "\\noindent{\\bfseries\\large\\color{accent}Skills}\\\\[4pt]\n"
        tex += ", ".join(skills) + "\\n\n\\vspace{6pt}\n"

    tech_skills = data.get("technical_skills", {})
    if tech_skills:
        lines = []
        for cat, items in tech_skills.items():
            if items:
                label = cat.replace("_", " ").title()
                lines.append(f"{label}: {', '.join(items)}")
        if lines:
            tex += "\\noindent{\\bfseries\\large\\color{accent}Technical Skills}\\\\[4pt]\n"
            for l in lines:
                tex += l + "\\\\\n"
            tex += "\\vspace{6pt}\n"

    experience = data.get("experience", [])
    if experience:
        tex += "\\noindent{\\bfseries\\large\\color{accent}Experience}\\\\[4pt]\n"
        for exp in experience:
            title = escape_latex(exp.get("title", ""))
            company = escape_latex(exp.get("company", ""))
            loc = escape_latex(exp.get("location", ""))
            sd = exp.get("start_date", "")
            ed = exp.get("end_date", "")
            tex += f"\\textbf{{{title}}} at \\textbf{{{company}}} \\hfill \\textit{{{sd} -- {ed}}}\\\\[2pt]\n"
            if loc:
                tex += f"\\textit{{{loc}}}\\\\[2pt]\n"
            bullets = exp.get("bullet_points", []) or [exp.get("description", "")]
            tex += "\\begin{itemize}[left=0pt]\n"
            for b in bullets:
                tex += "\\item " + escape_latex(b) + "\n"
            tex += "\\end{itemize}\n"
            tex += "\\vspace{3pt}\n"

    projects = data.get("projects", [])
    if projects:
        tex += "\\noindent{\\bfseries\\large\\color{accent}Projects}\\\\[4pt]\n"
        for proj in projects:
            pname = escape_latex(proj.get("name", ""))
            techs = proj.get("technologies", [])
            tex += f"\\textbf{{{pname}}}"
            if techs:
                tex += f" \\hfill {{{', '.join(techs)}}}\\\\[2pt]\n"
            else:
                tex += "\\\\[2pt]\n"
            desc = escape_latex(proj.get("description", ""))
            bullets = proj.get("bullet_points", [])
            items = bullets if bullets else ([desc] if desc else [])
            if items:
                tex += "\\begin{itemize}[left=0pt]\n"
                for b in items:
                    tex += "\\item " + escape_latex(b) + "\n"
                tex += "\\end{itemize}\n"
            tex += "\\vspace{3pt}\n"

    education = data.get("education", [])
    if education:
        tex += "\\noindent{\\bfseries\\large\\color{accent}Education}\\\\[4pt]\n"
        for edu in education:
            deg = escape_latex(edu.get("degree", ""))
            inst = escape_latex(edu.get("institution", ""))
            gd = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")
            tex += f"\\textbf{{{deg}}} -- {inst}"
            if gd:
                tex += f" \\hfill {gd}"
            tex += "\\\\[2pt]\n"
            if gpa:
                tex += f"\\textit{{GPA: {gpa}}}\\\\[2pt]\n"
            tex += "\\vspace{3pt}\n"

    certs = data.get("certifications", [])
    if certs:
        tex += "\\noindent{\\bfseries\\large\\color{accent}Certifications}\\\\[4pt]\n"
        tex += "\\begin{itemize}[left=0pt]\n"
        for c in certs:
            cname = escape_latex(c.get("name", ""))
            issuer = escape_latex(c.get("issuer", ""))
            cd = c.get("date", "")
            parts2 = filter(None, [cname, issuer, cd])
            tex += "\\item " + " -- ".join(parts2) + "\n"
        tex += "\\end{itemize}\n"

    achievements = data.get("achievements", [])
    if achievements:
        tex += "\\noindent{\\bfseries\\large\\color{accent}Achievements}\\\\[4pt]\n"
        tex += "\\begin{itemize}[left=0pt]\n"
        for a in achievements:
            tex += "\\item " + escape_latex(a) + "\n"
        tex += "\\end{itemize}\n"

    publications = data.get("publications", [])
    if publications:
        tex += "\\noindent{\\bfseries\\large\\color{accent}Publications}\\\\[4pt]\n"
        for pub in publications:
            ptitle = escape_latex(pub.get("title", ""))
            pjournal = escape_latex(pub.get("journal", ""))
            pdate = pub.get("date", "")
            plink = escape_latex(pub.get("link", ""))
            tex += "\\textbf{" + ptitle + "}"
            if pjournal:
                tex += " \\emph{" + pjournal + "}"
            if pdate:
                tex += " (" + pdate + ")"
            tex += "\\\\\n"
            if plink:
                tex += "\\url{" + plink + "}\\\\[4pt]\n"
            else:
                tex += "\\vspace{4pt}\n"

    tex += "\\end{document}\n"
    return tex


def build_latex_professional(data: dict) -> str:
    name = escape_latex(data.get("name", "Your Name"))
    email = escape_latex(data.get("email", ""))
    phone = escape_latex(data.get("phone", ""))
    location = escape_latex(data.get("location", ""))
    linkedin = escape_latex(data.get("linkedin", ""))
    github = escape_latex(data.get("github", ""))
    portfolio = escape_latex(data.get("portfolio", ""))
    summary = escape_latex(data.get("summary", ""))

    contact_parts = []
    for c in [email, phone, location]:
        if c:
            contact_parts.append(c)
    contact_line = " $|$ ".join(contact_parts)
    links = " $|$ ".join(filter(None, [linkedin, github, portfolio]))

    tex = "\\documentclass[10pt,a4paper]{article}\n"
    tex += "\\usepackage[utf8]{inputenc}\n"
    tex += "\\usepackage[T1]{fontenc}\n"
    tex += "\\usepackage{geometry}\n"
    tex += "\\geometry{margin=0.8in}\n"
    tex += "\\usepackage{parskip}\n"
    tex += "\\usepackage{hyperref}\n"
    tex += "\\hypersetup{colorlinks=true,linkcolor=black,urlcolor=black}\n"
    tex += "\\usepackage{enumitem}\n"
    tex += "\\setlist{nosep,left=0pt,label=\\textbullet,itemsep=1pt}\n"
    tex += "\\usepackage{xcolor}\n"
    tex += "\\definecolor{sectioncolor}{HTML}{1a1a2e}\n"
    tex += "\\pagestyle{empty}\n"
    tex += "\\begin{document}\n"
    tex += "\\begin{center}\n"
    tex += "{\\Huge\\bfseries\\color{sectioncolor} " + name + "}\\\\[6pt]\n"
    tex += contact_line + "\\\\\n"
    if links:
        tex += "\\small{" + links + "}\\\\\n"
    tex += "\\end{center}\n"
    tex += "\\vspace{-2pt}\n"
    tex += "{\\color{sectioncolor}\\rule{\\textwidth}{0.5pt}}\n"
    tex += "\\vspace{4pt}\n"

    if summary:
        tex += "\\noindent{\\bfseries\\large\\color{sectioncolor}Professional Summary}\\\\[4pt]\n"
        tex += summary + "\\n\n\\vspace{6pt}\n"

    skills = data.get("skills", [])
    if skills:
        tex += "\\noindent{\\bfseries\\large\\color{sectioncolor}Skills}\\\\[4pt]\n"
        tex += ", ".join(skills) + "\\n\n\\vspace{6pt}\n"

    tech_skills = data.get("technical_skills", {})
    if tech_skills:
        lines = []
        for cat, items in tech_skills.items():
            if items:
                label = cat.replace("_", " ").title()
                lines.append(f"\\textbf{{{label}}}: {', '.join(items)}")
        if lines:
            tex += "\\noindent{\\bfseries\\large\\color{sectioncolor}Technical Skills}\\\\[4pt]\n"
            for l in lines:
                tex += l + "\\\\\n"
            tex += "\\vspace{6pt}\n"

    experience = data.get("experience", [])
    if experience:
        tex += "\\noindent{\\bfseries\\large\\color{sectioncolor}Experience}\\\\[4pt]\n"
        for exp in experience:
            title = escape_latex(exp.get("title", ""))
            company = escape_latex(exp.get("company", ""))
            loc = escape_latex(exp.get("location", ""))
            sd = exp.get("start_date", "")
            ed = exp.get("end_date", "")
            tex += f"\\textbf{{{title}}} at \\textbf{{{company}}} \\hfill {sd} -- {ed}\\\\[2pt]\n"
            if loc:
                tex += f"\\textit{{{loc}}}\\\\[2pt]\n"
            bullets = exp.get("bullet_points", []) or [exp.get("description", "")]
            tex += "\\begin{itemize}[left=0pt]\n"
            for b in bullets:
                tex += "\\item " + escape_latex(b) + "\n"
            tex += "\\end{itemize}\n"
            tex += "\\vspace{2pt}\n"

    projects = data.get("projects", [])
    if projects:
        tex += "\\noindent{\\bfseries\\large\\color{sectioncolor}Projects}\\\\[4pt]\n"
        for proj in projects:
            pname = escape_latex(proj.get("name", ""))
            techs = proj.get("technologies", [])
            tex += f"\\textbf{{{pname}}}"
            if techs:
                tex += f" \\hfill {{{', '.join(techs)}}}\\\\[2pt]\n"
            else:
                tex += "\\\\[2pt]\n"
            desc = escape_latex(proj.get("description", ""))
            bullets = proj.get("bullet_points", [])
            items = bullets if bullets else ([desc] if desc else [])
            if items:
                tex += "\\begin{itemize}[left=0pt]\n"
                for b in items:
                    tex += "\\item " + escape_latex(b) + "\n"
                tex += "\\end{itemize}\n"
            tex += "\\vspace{2pt}\n"

    education = data.get("education", [])
    if education:
        tex += "\\noindent{\\bfseries\\large\\color{sectioncolor}Education}\\\\[4pt]\n"
        for edu in education:
            deg = escape_latex(edu.get("degree", ""))
            inst = escape_latex(edu.get("institution", ""))
            gd = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")
            tex += f"\\textbf{{{deg}}} -- {inst}"
            if gd:
                tex += f" \\hfill {gd}"
            tex += "\\\\[2pt]\n"
            if gpa:
                tex += f"GPA: {gpa}\\\\[2pt]\n"
            tex += "\\vspace{2pt}\n"

    certs = data.get("certifications", [])
    if certs:
        tex += "\\noindent{\\bfseries\\large\\color{sectioncolor}Certifications}\\\\[4pt]\n"
        tex += "\\begin{itemize}[left=0pt]\n"
        for c in certs:
            cname = escape_latex(c.get("name", ""))
            issuer = escape_latex(c.get("issuer", ""))
            cd = c.get("date", "")
            parts2 = filter(None, [cname, issuer, cd])
            tex += "\\item " + " -- ".join(parts2) + "\n"
        tex += "\\end{itemize}\n"

    achievements = data.get("achievements", [])
    if achievements:
        tex += "\\noindent{\\bfseries\\large\\color{sectioncolor}Achievements}\\\\[4pt]\n"
        tex += "\\begin{itemize}[left=0pt]\n"
        for a in achievements:
            tex += "\\item " + escape_latex(a) + "\n"
        tex += "\\end{itemize}\n"

    publications = data.get("publications", [])
    if publications:
        tex += "\\noindent{\\bfseries\\large\\color{sectioncolor}Publications}\\\\[4pt]\n"
        for pub in publications:
            ptitle = escape_latex(pub.get("title", ""))
            pjournal = escape_latex(pub.get("journal", ""))
            pdate = pub.get("date", "")
            plink = escape_latex(pub.get("link", ""))
            tex += "\\textbf{" + ptitle + "}"
            if pjournal:
                tex += " \\emph{" + pjournal + "}"
            if pdate:
                tex += " (" + pdate + ")"
            tex += "\\\\\n"
            if plink:
                tex += "\\url{" + plink + "}\\\\[4pt]\n"
            else:
                tex += "\\vspace{4pt}\n"

    tex += "\\end{document}\n"
    return tex


def build_latex_minimal(data: dict) -> str:
    name = escape_latex(data.get("name", "Your Name"))
    email = escape_latex(data.get("email", ""))
    phone = escape_latex(data.get("phone", ""))
    location = escape_latex(data.get("location", ""))
    linkedin = escape_latex(data.get("linkedin", ""))
    github = escape_latex(data.get("github", ""))
    portfolio = escape_latex(data.get("portfolio", ""))
    summary = escape_latex(data.get("summary", ""))

    contact_parts = []
    for c in [email, phone, location]:
        if c:
            contact_parts.append(c)
    contact_line = "  $|$  ".join(contact_parts)
    links = "  $|$  ".join(filter(None, [linkedin, github, portfolio]))

    tex = "\\documentclass[10pt,a4paper]{article}\n"
    tex += "\\usepackage[utf8]{inputenc}\n"
    tex += "\\usepackage[T1]{fontenc}\n"
    tex += "\\usepackage{geometry}\n"
    tex += "\\geometry{margin=0.9in}\n"
    tex += "\\usepackage{parskip}\n"
    tex += "\\usepackage{hyperref}\n"
    tex += "\\hypersetup{colorlinks=true,linkcolor=black,urlcolor=black}\n"
    tex += "\\usepackage{enumitem}\n"
    tex += "\\setlist{nosep,left=0pt,label=\\textbullet,itemsep=2pt}\n"
    tex += "\\pagestyle{empty}\n"
    tex += "\\begin{document}\n"
    tex += "\\begin{center}\n"
    tex += "{\\LARGE\\bfseries " + name + "}\\\\[4pt]\n"
    tex += "\\footnotesize{" + contact_line + "}\\\\\n"
    if links:
        tex += "\\footnotesize{" + links + "}\\\\\n"
    tex += "\\end{center}\n"
    tex += "\\vspace{6pt}\n"

    if summary:
        tex += "\\noindent{\\bfseries\\small " + summary + "}\\\\[8pt]\n"

    sections = []
    skills = data.get("skills", [])
    if skills:
        sections.append(("Skills", ", ".join(skills)))

    tech_skills = data.get("technical_skills", {})
    if tech_skills:
        lines = []
        for cat, items in tech_skills.items():
            if items:
                label = cat.replace("_", " ").title()
                lines.append(f"{label}: {', '.join(items)}")
        if lines:
            sections.append(("Technical Skills", " \\\\ ".join(lines)))

    experience = data.get("experience", [])
    if experience:
        exp_text = ""
        for exp in experience:
            title = escape_latex(exp.get("title", ""))
            company = escape_latex(exp.get("company", ""))
            loc = escape_latex(exp.get("location", ""))
            sd = exp.get("start_date", "")
            ed = exp.get("end_date", "")
            exp_text += f"\\textbf{{{title}}}, {company}"
            if loc:
                exp_text += f", {loc}"
            exp_text += f" \\hfill {sd} -- {ed}\\\\[2pt]\n"
            bullets = exp.get("bullet_points", []) or [exp.get("description", "")]
            for b in bullets:
                exp_text += "\\textbullet\\ " + escape_latex(b) + "\\\\\n"
            exp_text += "\\vspace{2pt}\n"
        sections.append(("Experience", exp_text))

    projects = data.get("projects", [])
    if projects:
        proj_text = ""
        for proj in projects:
            pname = escape_latex(proj.get("name", ""))
            techs = proj.get("technologies", [])
            proj_text += f"\\textbf{{{pname}}}"
            if techs:
                proj_text += f" \\hfill {{{', '.join(techs)}}}\\\\[2pt]\n"
            else:
                proj_text += "\\\\[2pt]\n"
            desc = escape_latex(proj.get("description", ""))
            bullets = proj.get("bullet_points", [])
            items = bullets if bullets else ([desc] if desc else [])
            for b in items:
                proj_text += "\\textbullet\\ " + escape_latex(b) + "\\\\\n"
            proj_text += "\\vspace{2pt}\n"
        sections.append(("Projects", proj_text))

    education = data.get("education", [])
    if education:
        edu_text = ""
        for edu in education:
            deg = escape_latex(edu.get("degree", ""))
            inst = escape_latex(edu.get("institution", ""))
            gd = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")
            edu_text += f"\\textbf{{{deg}}}, {inst}"
            if gd:
                edu_text += f" \\hfill {gd}"
            edu_text += "\\\\[2pt]\n"
            if gpa:
                edu_text += f"GPA: {gpa}\\\\[2pt]\n"
            edu_text += "\\vspace{2pt}\n"
        sections.append(("Education", edu_text))

    certs = data.get("certifications", [])
    if certs:
        cert_text = ""
        for c in certs:
            cname = escape_latex(c.get("name", ""))
            issuer = escape_latex(c.get("issuer", ""))
            cd = c.get("date", "")
            parts2 = filter(None, [cname, issuer, cd])
            cert_text += "\\textbullet\\ " + " -- ".join(parts2) + "\\\\\n"
        sections.append(("Certifications", cert_text))

    achievements = data.get("achievements", [])
    if achievements:
        ach_text = ""
        for a in achievements:
            ach_text += "\\textbullet\\ " + escape_latex(a) + "\\\\\n"
        sections.append(("Achievements", ach_text))

    publications = data.get("publications", [])
    if publications:
        pub_text = ""
        for pub in publications:
            ptitle = escape_latex(pub.get("title", ""))
            pjournal = escape_latex(pub.get("journal", ""))
            pdate = pub.get("date", "")
            plink = escape_latex(pub.get("link", ""))
            pub_text += "\\textbf{" + ptitle + "}"
            if pjournal:
                pub_text += " \\emph{" + pjournal + "}"
            if pdate:
                pub_text += " (" + pdate + ")"
            pub_text += "\\\\\n"
            if plink:
                pub_text += "\\url{" + plink + "}\\\\[4pt]\n"
            else:
                pub_text += "\\vspace{4pt}\n"
        sections.append(("Publications", pub_text))

    for i, (sec_title, sec_content) in enumerate(sections):
        tex += "\\noindent{\\bfseries\\small " + sec_title + "}\\\\[2pt]\n"
        tex += "\\rule{\\textwidth}{0.3pt}\\\\[2pt]\n"
        tex += sec_content + "\n"
        if i < len(sections) - 1:
            tex += "\\vspace{6pt}\n"

    tex += "\\end{document}\n"
    return tex


LATEX_BUILDERS = {
    "classic": build_latex_classic,
    "modern": build_latex_modern,
    "professional": build_latex_professional,
    "minimal": build_latex_minimal,
}


@router.post("/generate-latex")
async def generate_latex(resume_json: str = Form(...), template: str = Form("classic")):
    try:
        data = json.loads(resume_json)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid resume JSON")

    builder = LATEX_BUILDERS.get(template)
    if not builder:
        raise HTTPException(status_code=400, detail=f"Unknown template: {template}. Available: classic, modern, professional, minimal")

    tex = builder(data)
    return {"latex": tex}


def build_pdf_classic(pdf, data: dict):
    pdf.set_margins(15, 15, 15)

    def write_section(title, content):
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(0, 0, 0)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        if content:
            pdf.set_font("Helvetica", "", 9)
            if isinstance(content, str):
                pdf.multi_cell(0, 4.5, content)
            elif isinstance(content, list):
                for item in content:
                    pdf.cell(5, 4.5, "-")
                    pdf.multi_cell(0, 4.5, f" {item}")
        pdf.ln(3)

    def ct(t):
        return re.sub(r'\s+', ' ', str(t)).strip()

    name = data.get("name", "Your Name")
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 8, ct(name), new_x="LMARGIN", new_y="NEXT", align="C")

    contact_parts = []
    for field in ["email", "phone", "location"]:
        v = data.get(field, "")
        if v:
            contact_parts.append(ct(v))
    if contact_parts:
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(0, 5, " | ".join(contact_parts), new_x="LMARGIN", new_y="NEXT", align="C")

    link_parts = []
    for field in ["linkedin", "github", "portfolio"]:
        v = data.get(field, "")
        if v:
            link_parts.append(ct(v))
    if link_parts:
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(0, 5, " | ".join(link_parts), new_x="LMARGIN", new_y="NEXT", align="C")

    pdf.ln(2)
    pdf.set_draw_color(0, 0, 0)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(3)

    summary = data.get("summary", "")
    if summary:
        write_section("Professional Summary", ct(summary))
    skills = data.get("skills", [])
    if skills:
        write_section("Skills", [ct(s) for s in skills])
    tech_skills = data.get("technical_skills", {})
    if tech_skills:
        lines = []
        for cat, items in tech_skills.items():
            if items:
                label = cat.replace("_", " ").title()
                lines.append(f"{label}: {', '.join(ct(s) for s in items)}")
        if lines:
            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(0, 6, "Technical Skills", new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(0, 0, 0)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(2)
            pdf.set_font("Helvetica", "", 9)
            for l in lines:
                pdf.multi_cell(0, 4.5, l)
            pdf.ln(3)
    experience = data.get("experience", [])
    if experience:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "Experience", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(0, 0, 0)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        for exp in experience:
            title = ct(exp.get("title", ""))
            company = ct(exp.get("company", ""))
            loc = ct(exp.get("location", ""))
            sd = exp.get("start_date", "")
            ed = exp.get("end_date", "")
            pdf.set_font("Helvetica", "B", 9)
            date_str = f"{sd} -- {ed}" if sd or ed else ""
            loc_str = f" | {loc}" if loc else ""
            pdf.cell(0, 5, f"{title} at {company}{loc_str}", new_x="LMARGIN", new_y="NEXT")
            if date_str:
                pdf.set_font("Helvetica", "", 8)
                pdf.cell(0, 4, date_str, new_x="LMARGIN", new_y="NEXT")
            bullets = exp.get("bullet_points", []) or [exp.get("description", "")]
            pdf.set_font("Helvetica", "", 9)
            for b in bullets:
                bt = ct(b)
                if bt:
                    pdf.cell(5, 4.5, "-")
                    pdf.multi_cell(0, 4.5, f" {bt}")
            pdf.ln(2)
    projects = data.get("projects", [])
    if projects:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "Projects", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(0, 0, 0)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        for proj in projects:
            pname = ct(proj.get("name", ""))
            techs = proj.get("technologies", [])
            pdf.set_font("Helvetica", "B", 9)
            tech_str = f" | {', '.join(ct(t) for t in techs)}" if techs else ""
            pdf.cell(0, 5, f"{pname}{tech_str}", new_x="LMARGIN", new_y="NEXT")
            desc = ct(proj.get("description", ""))
            bullets = proj.get("bullet_points", [])
            items = bullets if bullets else ([desc] if desc else [])
            pdf.set_font("Helvetica", "", 9)
            for b in items:
                bt = ct(b)
                if bt:
                    pdf.cell(5, 4.5, "-")
                    pdf.multi_cell(0, 4.5, f" {bt}")
            pdf.ln(2)
    education = data.get("education", [])
    if education:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "Education", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(0, 0, 0)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        for edu in education:
            deg = ct(edu.get("degree", ""))
            inst = ct(edu.get("institution", ""))
            gd = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")
            pdf.set_font("Helvetica", "B", 9)
            gd_str = f" | {gd}" if gd else ""
            pdf.cell(0, 5, f"{deg} -- {inst}{gd_str}", new_x="LMARGIN", new_y="NEXT")
            if gpa:
                pdf.set_font("Helvetica", "", 8)
                pdf.cell(0, 4, f"GPA: {gpa}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
    certs = data.get("certifications", [])
    if certs:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "Certifications", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(0, 0, 0)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 9)
        for c in certs:
            parts2 = filter(None, [ct(c.get("name", "")), ct(c.get("issuer", "")), c.get("date", "")])
            pdf.cell(5, 4.5, "-")
            pdf.multi_cell(0, 4.5, f" {' -- '.join(parts2)}")
        pdf.ln(3)
    achievements = data.get("achievements", [])
    if achievements:
        write_section("Achievements", [ct(a) for a in achievements])
    publications = data.get("publications", [])
    if publications:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "Publications", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(0, 0, 0)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        for pub in publications:
            ptitle = ct(pub.get("title", ""))
            pjournal = ct(pub.get("journal", ""))
            pdate = pub.get("date", "")
            plink = ct(pub.get("link", ""))
            pdf.set_font("Helvetica", "B", 9)
            line = ptitle
            if pjournal:
                line += f" - {pjournal}"
            if pdate:
                line += f" ({pdate})"
            pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
            if plink:
                pdf.set_font("Helvetica", "", 8)
                pdf.cell(0, 4, plink, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)


def build_pdf_modern(pdf, data: dict):
    import struct
    pdf.set_margins(12, 12, 12)

    def ct(t):
        return re.sub(r'\s+', ' ', str(t)).strip()

    name = data.get("name", "Your Name")
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(14, 116, 144)
    pdf.cell(0, 10, ct(name), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_text_color(0, 0, 0)

    contact_parts = []
    for field in ["email", "phone", "location"]:
        v = data.get(field, "")
        if v:
            contact_parts.append(ct(v))
    if contact_parts:
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(0, 5, "  ".join(contact_parts), new_x="LMARGIN", new_y="NEXT", align="C")

    link_parts = []
    for field in ["linkedin", "github", "portfolio"]:
        v = data.get(field, "")
        if v:
            link_parts.append(ct(v))
    if link_parts:
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(14, 116, 144)
        pdf.cell(0, 5, "  ".join(link_parts), new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.set_text_color(0, 0, 0)

    pdf.ln(2)
    pdf.set_draw_color(14, 116, 144)
    pdf.set_line_width(0.8)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.set_line_width(0.2)
    pdf.ln(4)

    def write_section(title, content):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(14, 116, 144)
        pdf.cell(0, 6, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        if content:
            pdf.set_font("Helvetica", "", 9)
            if isinstance(content, str):
                pdf.multi_cell(0, 4.5, content)
            elif isinstance(content, list):
                for item in content:
                    pdf.cell(5, 4.5, chr(8226))
                    pdf.multi_cell(0, 4.5, f" {item}")
        pdf.ln(3)

    summary = data.get("summary", "")
    if summary:
        write_section("Professional Summary", ct(summary))
    skills = data.get("skills", [])
    if skills:
        write_section("Skills", [ct(s) for s in skills])
    tech_skills = data.get("technical_skills", {})
    if tech_skills:
        lines = []
        for cat, items in tech_skills.items():
            if items:
                label = cat.replace("_", " ").title()
                lines.append(f"{label}: {', '.join(ct(s) for s in items)}")
        if lines:
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(14, 116, 144)
            pdf.cell(0, 6, "Technical Skills", new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
            pdf.ln(1)
            pdf.set_font("Helvetica", "", 9)
            for l in lines:
                pdf.multi_cell(0, 4.5, l)
            pdf.ln(3)
    experience = data.get("experience", [])
    if experience:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(14, 116, 144)
        pdf.cell(0, 6, "Experience", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        for exp in experience:
            title = ct(exp.get("title", ""))
            company = ct(exp.get("company", ""))
            loc = ct(exp.get("location", ""))
            sd = exp.get("start_date", "")
            ed = exp.get("end_date", "")
            pdf.set_font("Helvetica", "B", 9)
            date_str = f"{sd} -- {ed}" if sd or ed else ""
            pdf.cell(0, 5, f"{title} at {company}", new_x="LMARGIN", new_y="NEXT")
            if loc or date_str:
                pdf.set_font("Helvetica", "", 8)
                parts = filter(None, [loc, date_str])
                pdf.cell(0, 4, " | ".join(parts), new_x="LMARGIN", new_y="NEXT")
            bullets = exp.get("bullet_points", []) or [exp.get("description", "")]
            pdf.set_font("Helvetica", "", 9)
            for b in bullets:
                bt = ct(b)
                if bt:
                    pdf.cell(5, 4.5, chr(8226))
                    pdf.multi_cell(0, 4.5, f" {bt}")
            pdf.ln(2)
    projects = data.get("projects", [])
    if projects:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(14, 116, 144)
        pdf.cell(0, 6, "Projects", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        for proj in projects:
            pname = ct(proj.get("name", ""))
            techs = proj.get("technologies", [])
            pdf.set_font("Helvetica", "B", 9)
            tech_str = f" | {', '.join(ct(t) for t in techs)}" if techs else ""
            pdf.cell(0, 5, f"{pname}{tech_str}", new_x="LMARGIN", new_y="NEXT")
            desc = ct(proj.get("description", ""))
            bullets = proj.get("bullet_points", [])
            items = bullets if bullets else ([desc] if desc else [])
            pdf.set_font("Helvetica", "", 9)
            for b in items:
                bt = ct(b)
                if bt:
                    pdf.cell(5, 4.5, chr(8226))
                    pdf.multi_cell(0, 4.5, f" {bt}")
            pdf.ln(2)
    education = data.get("education", [])
    if education:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(14, 116, 144)
        pdf.cell(0, 6, "Education", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        for edu in education:
            deg = ct(edu.get("degree", ""))
            inst = ct(edu.get("institution", ""))
            gd = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")
            pdf.set_font("Helvetica", "B", 9)
            gd_str = f" | {gd}" if gd else ""
            pdf.cell(0, 5, f"{deg} -- {inst}{gd_str}", new_x="LMARGIN", new_y="NEXT")
            if gpa:
                pdf.set_font("Helvetica", "", 8)
                pdf.cell(0, 4, f"GPA: {gpa}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
    certs = data.get("certifications", [])
    if certs:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(14, 116, 144)
        pdf.cell(0, 6, "Certifications", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        pdf.set_font("Helvetica", "", 9)
        for c in certs:
            parts2 = filter(None, [ct(c.get("name", "")), ct(c.get("issuer", "")), c.get("date", "")])
            pdf.cell(5, 4.5, chr(8226))
            pdf.multi_cell(0, 4.5, f" {' -- '.join(parts2)}")
        pdf.ln(3)
    achievements = data.get("achievements", [])
    if achievements:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(14, 116, 144)
        pdf.cell(0, 6, "Achievements", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        pdf.set_font("Helvetica", "", 9)
        for a in achievements:
            pdf.cell(5, 4.5, chr(8226))
            pdf.multi_cell(0, 4.5, f" {ct(a)}")
        pdf.ln(3)
    publications = data.get("publications", [])
    if publications:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(14, 116, 144)
        pdf.cell(0, 6, "Publications", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        for pub in publications:
            ptitle = ct(pub.get("title", ""))
            pjournal = ct(pub.get("journal", ""))
            pdate = pub.get("date", "")
            pdf.set_font("Helvetica", "B", 9)
            line = ptitle
            if pjournal:
                line += f" - {pjournal}"
            if pdate:
                line += f" ({pdate})"
            pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)


def build_pdf_professional(pdf, data: dict):
    pdf.set_margins(18, 18, 18)

    def ct(t):
        return re.sub(r'\s+', ' ', str(t)).strip()

    name = data.get("name", "Your Name")
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(26, 26, 46)
    pdf.cell(0, 8, ct(name), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_text_color(0, 0, 0)

    contact_parts = []
    for field in ["email", "phone", "location"]:
        v = data.get(field, "")
        if v:
            contact_parts.append(ct(v))
    if contact_parts:
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(0, 5, " | ".join(contact_parts), new_x="LMARGIN", new_y="NEXT", align="C")
    link_parts = []
    for field in ["linkedin", "github", "portfolio"]:
        v = data.get(field, "")
        if v:
            link_parts.append(ct(v))
    if link_parts:
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(0, 5, " | ".join(link_parts), new_x="LMARGIN", new_y="NEXT", align="C")

    pdf.ln(2)
    pdf.set_draw_color(26, 26, 46)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(4)

    def write_section(title, content):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 6, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        if content:
            pdf.set_font("Helvetica", "", 9)
            if isinstance(content, str):
                pdf.multi_cell(0, 4.5, content)
            elif isinstance(content, list):
                for item in content:
                    pdf.cell(5, 4.5, "-")
                    pdf.multi_cell(0, 4.5, f" {item}")
        pdf.ln(3)

    summary = data.get("summary", "")
    if summary:
        write_section("Professional Summary", ct(summary))
    skills = data.get("skills", [])
    if skills:
        write_section("Skills", [ct(s) for s in skills])
    tech_skills = data.get("technical_skills", {})
    if tech_skills:
        lines = []
        for cat, items in tech_skills.items():
            if items:
                label = cat.replace("_", " ").title()
                lines.append(f"{label}: {', '.join(ct(s) for s in items)}")
        if lines:
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(26, 26, 46)
            pdf.cell(0, 6, "Technical Skills", new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
            pdf.ln(1)
            pdf.set_font("Helvetica", "", 9)
            for l in lines:
                pdf.multi_cell(0, 4.5, l)
            pdf.ln(3)
    experience = data.get("experience", [])
    if experience:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 6, "Experience", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        for exp in experience:
            title = ct(exp.get("title", ""))
            company = ct(exp.get("company", ""))
            loc = ct(exp.get("location", ""))
            sd = exp.get("start_date", "")
            ed = exp.get("end_date", "")
            pdf.set_font("Helvetica", "B", 9)
            date_str = f"{sd} -- {ed}" if sd or ed else ""
            loc_str = f" | {loc}" if loc else ""
            pdf.cell(0, 5, f"{title} at {company}{loc_str}", new_x="LMARGIN", new_y="NEXT")
            if date_str:
                pdf.set_font("Helvetica", "", 8)
                pdf.cell(0, 4, date_str, new_x="LMARGIN", new_y="NEXT")
            bullets = exp.get("bullet_points", []) or [exp.get("description", "")]
            pdf.set_font("Helvetica", "", 9)
            for b in bullets:
                bt = ct(b)
                if bt:
                    pdf.cell(5, 4.5, "-")
                    pdf.multi_cell(0, 4.5, f" {bt}")
            pdf.ln(2)
    projects = data.get("projects", [])
    if projects:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 6, "Projects", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        for proj in projects:
            pname = ct(proj.get("name", ""))
            techs = proj.get("technologies", [])
            pdf.set_font("Helvetica", "B", 9)
            tech_str = f" | {', '.join(ct(t) for t in techs)}" if techs else ""
            pdf.cell(0, 5, f"{pname}{tech_str}", new_x="LMARGIN", new_y="NEXT")
            desc = ct(proj.get("description", ""))
            bullets = proj.get("bullet_points", [])
            items = bullets if bullets else ([desc] if desc else [])
            pdf.set_font("Helvetica", "", 9)
            for b in items:
                bt = ct(b)
                if bt:
                    pdf.cell(5, 4.5, "-")
                    pdf.multi_cell(0, 4.5, f" {bt}")
            pdf.ln(2)
    education = data.get("education", [])
    if education:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 6, "Education", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        for edu in education:
            deg = ct(edu.get("degree", ""))
            inst = ct(edu.get("institution", ""))
            gd = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")
            pdf.set_font("Helvetica", "B", 9)
            gd_str = f" | {gd}" if gd else ""
            pdf.cell(0, 5, f"{deg} -- {inst}{gd_str}", new_x="LMARGIN", new_y="NEXT")
            if gpa:
                pdf.set_font("Helvetica", "", 8)
                pdf.cell(0, 4, f"GPA: {gpa}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
    certs = data.get("certifications", [])
    if certs:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 6, "Certifications", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        pdf.set_font("Helvetica", "", 9)
        for c in certs:
            parts2 = filter(None, [ct(c.get("name", "")), ct(c.get("issuer", "")), c.get("date", "")])
            pdf.cell(5, 4.5, "-")
            pdf.multi_cell(0, 4.5, f" {' -- '.join(parts2)}")
        pdf.ln(3)
    achievements = data.get("achievements", [])
    if achievements:
        write_section("Achievements", [ct(a) for a in achievements])
    publications = data.get("publications", [])
    if publications:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 6, "Publications", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        for pub in publications:
            ptitle = ct(pub.get("title", ""))
            pjournal = ct(pub.get("journal", ""))
            pdate = pub.get("date", "")
            plink = ct(pub.get("link", ""))
            pdf.set_font("Helvetica", "B", 9)
            line = ptitle
            if pjournal:
                line += f" - {pjournal}"
            if pdate:
                line += f" ({pdate})"
            pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
            if plink:
                pdf.set_font("Helvetica", "", 8)
                pdf.cell(0, 4, plink, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)


def build_pdf_minimal(pdf, data: dict):
    pdf.set_margins(20, 20, 20)

    def ct(t):
        return re.sub(r'\s+', ' ', str(t)).strip()

    name = data.get("name", "Your Name")
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 8, ct(name), new_x="LMARGIN", new_y="NEXT", align="C")

    contact_parts = []
    for field in ["email", "phone", "location"]:
        v = data.get(field, "")
        if v:
            contact_parts.append(ct(v))
    if contact_parts:
        pdf.set_font("Helvetica", "", 7)
        pdf.cell(0, 4, "  |  ".join(contact_parts), new_x="LMARGIN", new_y="NEXT", align="C")
    link_parts = []
    for field in ["linkedin", "github", "portfolio"]:
        v = data.get(field, "")
        if v:
            link_parts.append(ct(v))
    if link_parts:
        pdf.set_font("Helvetica", "", 7)
        pdf.cell(0, 4, "  |  ".join(link_parts), new_x="LMARGIN", new_y="NEXT", align="C")

    pdf.ln(4)

    def write_section(title, content):
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(0, 5, title.upper(), new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(180, 180, 180)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        pdf.set_draw_color(0, 0, 0)
        if content:
            pdf.set_font("Helvetica", "", 8)
            if isinstance(content, str):
                pdf.multi_cell(0, 4, content)
            elif isinstance(content, list):
                for item in content:
                    pdf.cell(3, 4, "-")
                    pdf.multi_cell(0, 4, f" {item}")
        pdf.ln(2)

    summary = data.get("summary", "")
    if summary:
        pdf.set_font("Helvetica", "", 8)
        pdf.multi_cell(0, 4, ct(summary))
        pdf.ln(4)

    skills = data.get("skills", [])
    if skills:
        write_section("Skills", [ct(s) for s in skills])

    tech_skills = data.get("technical_skills", {})
    if tech_skills:
        lines = []
        for cat, items in tech_skills.items():
            if items:
                label = cat.replace("_", " ").title()
                lines.append(f"{label}: {', '.join(ct(s) for s in items)}")
        if lines:
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 5, "TECHNICAL SKILLS", new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(180, 180, 180)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(2)
            pdf.set_draw_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 8)
            for l in lines:
                pdf.multi_cell(0, 4, l)
            pdf.ln(2)

    experience = data.get("experience", [])
    if experience:
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(0, 5, "EXPERIENCE", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(180, 180, 180)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        pdf.set_draw_color(0, 0, 0)
        for exp in experience:
            title = ct(exp.get("title", ""))
            company = ct(exp.get("company", ""))
            loc = ct(exp.get("location", ""))
            sd = exp.get("start_date", "")
            ed = exp.get("end_date", "")
            pdf.set_font("Helvetica", "B", 8)
            date_str = f"{sd} -- {ed}" if sd or ed else ""
            pdf.cell(0, 4, f"{title}, {company}", new_x="LMARGIN", new_y="NEXT")
            if loc or date_str:
                pdf.set_font("Helvetica", "", 7)
                parts = filter(None, [loc, date_str])
                pdf.cell(0, 3, ", ".join(parts), new_x="LMARGIN", new_y="NEXT")
            bullets = exp.get("bullet_points", []) or [exp.get("description", "")]
            pdf.set_font("Helvetica", "", 8)
            for b in bullets:
                bt = ct(b)
                if bt:
                    pdf.cell(3, 3.5, "-")
                    pdf.multi_cell(0, 3.5, f" {bt}")
            pdf.ln(1)

    projects = data.get("projects", [])
    if projects:
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(0, 5, "PROJECTS", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(180, 180, 180)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        pdf.set_draw_color(0, 0, 0)
        for proj in projects:
            pname = ct(proj.get("name", ""))
            techs = proj.get("technologies", [])
            pdf.set_font("Helvetica", "B", 8)
            tech_str = f" | {', '.join(ct(t) for t in techs)}" if techs else ""
            pdf.cell(0, 4, f"{pname}{tech_str}", new_x="LMARGIN", new_y="NEXT")
            desc = ct(proj.get("description", ""))
            bullets = proj.get("bullet_points", [])
            items = bullets if bullets else ([desc] if desc else [])
            pdf.set_font("Helvetica", "", 8)
            for b in items:
                bt = ct(b)
                if bt:
                    pdf.cell(3, 3.5, "-")
                    pdf.multi_cell(0, 3.5, f" {bt}")
            pdf.ln(1)

    education = data.get("education", [])
    if education:
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(0, 5, "EDUCATION", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(180, 180, 180)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        pdf.set_draw_color(0, 0, 0)
        for edu in education:
            deg = ct(edu.get("degree", ""))
            inst = ct(edu.get("institution", ""))
            gd = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")
            pdf.set_font("Helvetica", "B", 8)
            gd_str = f" | {gd}" if gd else ""
            pdf.cell(0, 4, f"{deg} -- {inst}{gd_str}", new_x="LMARGIN", new_y="NEXT")
            if gpa:
                pdf.set_font("Helvetica", "", 7)
                pdf.cell(0, 3, f"GPA: {gpa}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

    certs = data.get("certifications", [])
    if certs:
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(0, 5, "CERTIFICATIONS", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(180, 180, 180)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        pdf.set_draw_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 8)
        for c in certs:
            parts2 = filter(None, [ct(c.get("name", "")), ct(c.get("issuer", "")), c.get("date", "")])
            pdf.cell(3, 3.5, "-")
            pdf.multi_cell(0, 3.5, f" {' -- '.join(parts2)}")
        pdf.ln(2)

    achievements = data.get("achievements", [])
    if achievements:
        write_section("Achievements", [ct(a) for a in achievements])

    publications = data.get("publications", [])
    if publications:
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(0, 5, "PUBLICATIONS", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(180, 180, 180)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        pdf.set_draw_color(0, 0, 0)
        for pub in publications:
            ptitle = ct(pub.get("title", ""))
            pjournal = ct(pub.get("journal", ""))
            pdate = pub.get("date", "")
            pdf.set_font("Helvetica", "B", 8)
            line = ptitle
            if pjournal:
                line += f" - {pjournal}"
            if pdate:
                line += f" ({pdate})"
            pdf.cell(0, 4, line, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)


PDF_BUILDERS = {
    "classic": build_pdf_classic,
    "modern": build_pdf_modern,
    "professional": build_pdf_professional,
    "minimal": build_pdf_minimal,
}


@router.post("/generate-pdf")
async def generate_pdf(resume_json: str = Form(...), template: str = Form("classic")):
    try:
        data = json.loads(resume_json)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid resume JSON")

    builder = PDF_BUILDERS.get(template)
    if not builder:
        raise HTTPException(status_code=400, detail=f"Unknown template: {template}. Available: classic, modern, professional, minimal")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    builder(pdf, data)

    pdf_bytes = pdf.output()
    return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=optimized_resume.pdf"})

@router.post("/generate-docx")
async def generate_docx(resume_json: str = Form(...), template: str = Form("classic")):
    try:
        data = json.loads(resume_json)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid resume JSON")
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        def add_heading_text(text, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
            p = doc.add_paragraph()
            p.alignment = align
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = 'Calibri'
            return p

        def add_separator():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): '000000',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)

        def add_section_title(text):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            add_separator()
            return p

        def add_bullet(text):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            return p

        def add_para(text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
            p = doc.add_paragraph()
            p.alignment = align
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.name = 'Calibri'
            run.bold = bold
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            return p

        name = data.get("name", "Your Name")
        add_heading_text(name, 16)

        contact_parts = []
        for field in ["email", "phone", "location"]:
            v = data.get(field, "")
            if v:
                contact_parts.append(v)
        if contact_parts:
            add_para(" | ".join(contact_parts), 8, align=WD_ALIGN_PARAGRAPH.CENTER)

        link_parts = []
        for field in ["linkedin", "github", "portfolio"]:
            v = data.get(field, "")
            if v:
                link_parts.append(v)
        if link_parts:
            add_para(" | ".join(link_parts), 8, align=WD_ALIGN_PARAGRAPH.CENTER)

        summary = data.get("summary", "")
        if summary:
            add_section_title("Professional Summary")
            add_para(summary, 9)

        skills = data.get("skills", [])
        if skills:
            add_section_title("Skills")
            add_para(", ".join(skills), 9)

        tech_skills = data.get("technical_skills", {})
        if tech_skills:
            add_section_title("Technical Skills")
            for cat, items in tech_skills.items():
                if items:
                    label = cat.replace("_", " ").title()
                    add_para(f"{label}: {', '.join(items)}", 9)

        experience = data.get("experience", [])
        if experience:
            add_section_title("Experience")
            for exp in experience:
                title = exp.get("title", "")
                company = exp.get("company", "")
                loc = exp.get("location", "")
                sd = exp.get("start_date", "")
                ed = exp.get("end_date", "")
                header = f"{title} at {company}"
                if loc:
                    header += f" | {loc}"
                add_para(header, 10, bold=True)
                if sd or ed:
                    add_para(f"{sd} -- {ed}", 8)
                bullets = exp.get("bullet_points", []) or [exp.get("description", "")]
                for b in bullets:
                    add_bullet(b)

        projects = data.get("projects", [])
        if projects:
            add_section_title("Projects")
            for proj in projects:
                pname = proj.get("name", "")
                techs = proj.get("technologies", [])
                header = pname
                if techs:
                    header += f" | {', '.join(techs)}"
                add_para(header, 10, bold=True)
                desc = proj.get("description", "")
                bullets = proj.get("bullet_points", [])
                items = bullets if bullets else ([desc] if desc else [])
                for b in items:
                    add_bullet(b)

        education = data.get("education", [])
        if education:
            add_section_title("Education")
            for edu in education:
                deg = edu.get("degree", "")
                inst = edu.get("institution", "")
                gd = edu.get("graduation_date", "")
                gpa = edu.get("gpa", "")
                header = f"{deg} -- {inst}"
                if gd:
                    header += f" | {gd}"
                add_para(header, 10, bold=True)
                if gpa:
                    add_para(f"GPA: {gpa}", 8)

        certs = data.get("certifications", [])
        if certs:
            add_section_title("Certifications")
            for c in certs:
                parts2 = filter(None, [c.get("name", ""), c.get("issuer", ""), c.get("date", "")])
                add_bullet(" -- ".join(parts2))

        achievements = data.get("achievements", [])
        if achievements:
            add_section_title("Achievements")
            for a in achievements:
                add_bullet(a)

        publications = data.get("publications", [])
        if publications:
            add_section_title("Publications")
            for pub in publications:
                parts3 = filter(None, [pub.get("title", ""), pub.get("journal", ""), pub.get("date", "")])
                add_bullet(" -- ".join(parts3))

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return Response(content=buf.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=optimized_resume.docx"})
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX generation library not available")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")

def section_block(title, content):
    tex = "\\section*{" + title + "}\n"
    tex += "\\vspace{-2pt}\\rule{\\textwidth}{0.5pt}\\vspace{4pt}\n"
    if isinstance(content, str):
        tex += content + "\n\n"
    elif isinstance(content, list):
        tex += "\\begin{itemize}[left=0pt]\n"
        for item in content:
            tex += "\\item " + item + "\n"
        tex += "\\end{itemize}\n"
    tex += "\\vspace{4pt}\n"
    return tex

def escape_latex(text):
    if not text:
        return ""
    replacements = {
        '\\': '\\textbackslash{}',
        '{': '\\{',
        '}': '\\}',
        '$': '\\$',
        '&': '\\&',
        '#': '\\#',
        '_': '\\_',
        '%': '\\%',
        '~': '\\textasciitilde{}',
        '^': '\\textasciicircum{}',
    }
    for k, v in replacements.items():
        if k == '\\':
            text = text.replace(k, v)
    for k, v in replacements.items():
        if k != '\\':
            text = text.replace(k, v)
    return text
